"""
Document management system for AlphaAnalyst Trading AI Agent
Enhanced for Phase 2 with PDF processing and AI analysis
"""
import os
import io
import uuid
from datetime import datetime
from typing import List, Dict, Optional, BinaryIO
import PyPDF2
import docx
import docx2txt
import numpy as np
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

"""Gemini-only embedding configuration (sentence-transformers disabled by design)."""

from ..database.config import get_supabase
from groq import Groq

class DocumentManager:
    """Document management for research documents"""
    
    def __init__(self):
        self.supabase = get_supabase()
        # Initialize sentence transformer for embeddings (if available)
        self.embedding_model = None  # Explicitly unused (Gemini-only embeddings)
        self.gemini_client = None
        # Gemini embedding provider
        try:
            import google.generativeai as genai
            gemini_key = os.getenv("GEMINI_API_KEY", "")
            if gemini_key:
                genai.configure(api_key=gemini_key)
                self.gemini_client = genai
        except Exception as e:
            self.gemini_client = None
    
    def _extract_text_from_stream(self, file_content: BinaryIO, file_extension: str) -> str:
        """Extract text content from file stream (no local storage)"""
        try:
            if file_extension.lower() == '.pdf':
                return self._extract_pdf_from_stream(file_content)
            elif file_extension.lower() in ['.docx', '.doc']:
                return self._extract_docx_from_stream(file_content)
            elif file_extension.lower() in ['.txt']:
                return self._extract_txt_from_stream(file_content)
            else:
                # Try alternative extraction for unknown formats
                return self._extract_with_alternative_method(file_content, file_extension)
        except Exception as e:
            print(f"Error extracting text from stream: {e}")
            return f"Error extracting text from uploaded file"
    
    def _extract_text_from_file(self, file_path: str, file_extension: str) -> str:
        """Extract text content from various file formats (legacy method)"""
        try:
            if file_extension.lower() == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension.lower() in ['.docx', '.doc']:
                return self._extract_docx_text(file_path)
            elif file_extension.lower() in ['.txt']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                # Try alternative extraction for unknown formats
                return self._extract_with_alternative_method_from_file(file_path, file_extension)
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return f"Error extracting text from {file_path}"
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
            return ""
    
    def _extract_pdf_from_stream(self, file_content: BinaryIO) -> str:
        """Extract text from PDF stream"""
        text = ""
        try:
            # Reset file pointer to beginning
            if hasattr(file_content, 'seek'):
                file_content.seek(0)
            
            # Read PDF from stream
            if hasattr(file_content, 'getbuffer'):
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content.getbuffer()))
            else:
                pdf_reader = PyPDF2.PdfReader(file_content)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error reading PDF from stream: {e}")
        return text
    
    def _extract_docx_from_stream(self, file_content: BinaryIO) -> str:
        """Extract text from DOCX stream"""
        try:
            # Reset file pointer to beginning
            if hasattr(file_content, 'seek'):
                file_content.seek(0)
            
            # Read DOCX from stream
            if hasattr(file_content, 'getbuffer'):
                doc = docx.Document(io.BytesIO(file_content.getbuffer()))
            else:
                doc = docx.Document(file_content)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error reading DOCX from stream: {e}")
            return ""
    
    def _extract_txt_from_stream(self, file_content: BinaryIO) -> str:
        """Extract text from TXT stream"""
        try:
            # Reset file pointer to beginning
            if hasattr(file_content, 'seek'):
                file_content.seek(0)
            
            # Read text from stream
            if hasattr(file_content, 'getbuffer'):
                return file_content.getbuffer().decode('utf-8')
            else:
                return file_content.read().decode('utf-8')
        except Exception as e:
            print(f"Error reading TXT from stream: {e}")
            return ""
    
    def _extract_with_alternative_method(self, file_content: BinaryIO, file_extension: str) -> str:
        """Extract text using alternative methods for unsupported formats"""
        try:
            # Reset file pointer to beginning
            if hasattr(file_content, 'seek'):
                file_content.seek(0)
            
            # For .doc files (old Word format), try to use text from file
            if file_extension.lower() == '.doc':
                # Try reading as text first
                try:
                    if hasattr(file_content, 'getbuffer'):
                        return file_content.getbuffer().decode('utf-8', errors='ignore')
                    else:
                        return file_content.read().decode('utf-8', errors='ignore')
                except:
                    pass
            
            # Create temporary file and try docx2txt
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                if hasattr(file_content, 'getbuffer'):
                    temp_file.write(file_content.getbuffer())
                else:
                    temp_file.write(file_content.read())
                temp_file.flush()
                
                try:
                    # Try docx2txt for .doc files
                    if file_extension.lower() in ['.doc', '.docx']:
                        result = docx2txt.process(temp_file.name)
                        if result:
                            return result
                except:
                    pass
                
                # Fallback: try reading as raw text
                try:
                    with open(temp_file.name, 'r', encoding='utf-8', errors='ignore') as f:
                        result = f.read()
                    if result:
                        return result
                except:
                    pass
                
                # Clean up temporary file
                os.unlink(temp_file.name)
                return f"Unable to extract text from {file_extension} format"
        except Exception as e:
            print(f"Error using alternative extraction on stream: {e}")
            return ""
    
    def _extract_with_alternative_method_from_file(self, file_path: str, file_extension: str) -> str:
        """Extract text using alternative methods from file path"""
        try:
            # For .doc files (old Word format)
            if file_extension.lower() == '.doc':
                try:
                    return docx2txt.process(file_path)
                except:
                    pass
            
            # Try reading as raw text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                result = f.read()
            if result:
                return result
            return f"Unable to extract text from {file_extension} format"
        except Exception as e:
            print(f"Error using alternative extraction from file: {e}")
            return f"Unable to extract text from {file_extension} format"
    
    def _generate_embedding(self, text: str) -> Optional[List[float]]:
        """Generate embedding vector for text"""
        if not text:
            return None
        # Gemini-only embeddings
        if self.gemini_client:
            try:
                model = self.gemini_client.GenerativeModel("text-embedding-004")
                resp = model.embed_content(text)
                vec = resp.get("embedding", {}).get("values") if isinstance(resp, dict) else getattr(resp, "embedding", None)
                if hasattr(vec, "values"):
                    return list(vec.values)
                if isinstance(vec, list):
                    return vec
            except Exception as e:
                print(f"Embedding (Gemini) failed: {e}")
        # No embedding available
        return None
    
    def upload_document(self, 
                       file_content: BinaryIO, 
                       filename: str, 
                       title: str,
                       document_type: str = "research",
                       symbol: Optional[str] = None,
                       source: str = "user_upload") -> Dict:
        """Upload and store a document (cloud-only storage)"""
        try:
            # File size validation (10MB limit)
            MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
            if hasattr(file_content, 'getbuffer'):
                file_size = len(file_content.getbuffer())
            elif hasattr(file_content, 'read'):
                file_content.seek(0, 2)  # Seek to end
                file_size = file_content.tell()
                file_content.seek(0)  # Reset to beginning
            else:
                file_size = len(file_content)
            
            if file_size > MAX_FILE_SIZE:
                return {"success": False, "error": "File too large", "message": f"File size {file_size/1024/1024:.1f}MB exceeds 10MB limit"}
            
            # Extract text content directly from file stream (no local storage)
            file_extension = os.path.splitext(filename)[1]
            content_text = self._extract_text_from_stream(file_content, file_extension)
            
            # Generate embedding for the content
            embedding_vector = self._generate_embedding(content_text)

            if self.supabase:
                # Insert minimal required columns first
                base_payload = {
                    "file_name": filename,
                    "file_content": content_text or f"Document: {filename}",
                    "symbol": symbol,
                    "uploaded_at": datetime.now().isoformat(),
                }
                try:
                    resp = self.supabase.table("research_documents").insert(base_payload).execute()
                    print(f"Supabase insert response: {resp}")
                except Exception as e:
                    print(f"Supabase insert error details: {e}")
                    return {"success": False, "error": str(e), "message": "Supabase insert failed"}

                # Try to update optional RAG fields if the columns exist
                try:
                    new_id = resp.data[0]['id'] if resp.data else None
                    if new_id:
                        optional_update = {
                            "rag_embedding": embedding_vector if embedding_vector is not None else None,
                            "doc_title": title,
                            "document_type": document_type,
                            "source": source
                        }
                        self.supabase.table("research_documents").update(optional_update).eq("id", new_id).execute()
                except Exception as e:
                    # Non-fatal: schema may not have optional columns
                    print(f"Supabase optional update skipped: {e}")

                return {"success": True, "message": "Document uploaded to cloud storage", "document_id": resp.data[0]['id'] if resp.data else None}
            else:
                return {"success": False, "error": "Supabase not configured", "message": "Cannot upload document"}
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to upload document"
            }
    
    def get_documents(self, symbol: Optional[str] = None, 
                     document_type: Optional[str] = None) -> List[Dict]:
        """Get documents with optional filtering"""
        try:
            if not self.supabase:
                print("Supabase not configured. Cannot retrieve documents.")
                return []
                
            query = self.supabase.table("research_documents").select("*")
            # Filter by symbol if provided
            if symbol:
                query = query.eq("symbol", symbol)
            
            resp = query.execute()
            rows = resp.data if hasattr(resp, "data") else []
            return rows
            
        except Exception as e:
            print(f"Error retrieving documents: {e}")
            return []
    
    def get_document_content(self, document_id: str) -> Optional[str]:
        """Get document content by ID"""
        try:
            if not self.supabase:
                print("Supabase not configured. Cannot retrieve document content.")
                return None
                
            resp = self.supabase.table("research_documents").select("file_content").eq("id", document_id).execute()
            data = resp.data if hasattr(resp, "data") else []
            if data and len(data) > 0:
                return data[0].get("file_content")
            return None
        except Exception as e:
            print(f"Error reading document {document_id}: {e}")
            return None
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document"""
        try:
            if not self.supabase:
                print("Supabase not configured. Cannot delete document.")
                return False
                
            # Delete from Supabase
            resp = self.supabase.table("research_documents").delete().eq("id", document_id).execute()
            return True
        except Exception as e:
            print(f"Error deleting document {document_id}: {e}")
            return False
    
    def get_document_stats(self) -> Dict:
        """Get document statistics"""
        try:
            if not self.supabase:
                print("Supabase not configured. Cannot get document stats.")
                return {"total_documents": 0, "by_type": {}, "by_source": {}}
                
            # Get all documents from Supabase
            resp = self.supabase.table("research_documents").select("*").execute()
            documents = resp.data if hasattr(resp, "data") else []
            
            total_docs = len(documents)
            docs_by_symbol = {}
            
            for doc in documents:
                metadata = doc.get("metadata", {})
                symbol = metadata.get("symbol", "Unknown")
                if symbol not in docs_by_symbol:
                    docs_by_symbol[symbol] = 0
                docs_by_symbol[symbol] += 1
            
            return {
                "total_documents": total_docs,
                "by_symbol": docs_by_symbol
            }
        except Exception as e:
            print(f"Error getting document stats: {e}")
            return {"total_documents": 0, "by_symbol": {}}
    
    def analyze_document_with_ai(self, document_id: str, symbol: str = None) -> Dict:
        """Analyze document content using AI to extract trading insights"""
        try:
            # Get document content
            content = self.get_document_content(document_id)
            if not content:
                return {"success": False, "error": "Document content not found"}
            
            # Use xAI/Phi for analysis
            from phi.model.xai import xAI
            from phi.agent.agent import Agent
            
            # Initialize xAI model only with its own key; fallback to Groq later
            xai_key = os.getenv("XAI_API_KEY", "")
            xai_model = xAI(model="grok-beta", api_key=xai_key) if xai_key else None
            
            # Create analysis prompt
            analysis_prompt = f"""
            Analyze the following research document for trading insights related to {symbol or 'the mentioned stock'}.
            
            Document Content:
            {content[:4000]}  # Limit content to avoid token limits
            
            Please provide:
            1. Key financial metrics mentioned
            2. Bullish signals (positive indicators)
            3. Bearish signals (negative indicators)
            4. Overall sentiment (Bullish/Bearish/Neutral)
            5. Confidence level (1-10)
            6. Key risks mentioned
            7. Investment recommendation (BUY/SELL/HOLD)
            
            Format your response as a structured analysis.
            """
            
            # Get AI analysis using Agent (xAI), fallback to Groq if needed
            response = None
            if xai_model:
                try:
                    agent = Agent(model=xai_model)
                    response = agent.run(analysis_prompt)
                except Exception:
                    response = None
            if response is None:
                groq_key = os.getenv("GROQ_API_KEY", "")
                if groq_key:
                    client = Groq(api_key=groq_key)
                    chat = client.chat.completions.create(
                        model="llama-3.1-70b-versatile",
                        messages=[
                            {"role": "system", "content": "You are a professional financial analyst."},
                            {"role": "user", "content": analysis_prompt},
                        ],
                        temperature=0.2,
                    )
                    response = chat.choices[0].message.content if chat.choices else ""
                else:
                    return {"success": False, "error": "No working provider. Set XAI_API_KEY or GROQ_API_KEY"}
            
            return {
                "success": True,
                "analysis": response,
                "document_id": document_id,
                "symbol": symbol,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            msg = str(e)
            if "403" in msg or "permission" in msg.lower() or "credits" in msg.lower():
                return {"success": False, "error": "Permission/credits issue with xAI. Set XAI_API_KEY and ensure team has credits.", "details": msg}
            return {"success": False, "error": msg}
    
    def extract_trading_signals(self, document_id: str) -> Dict:
        """Extract specific trading signals from document"""
        try:
            content = self.get_document_content(document_id)
            if not content:
                return {"success": False, "error": "Document content not found"}
            
            # Look for common trading signal keywords
            bullish_keywords = [
                "buy", "bullish", "positive", "growth", "outperform", "upgrade",
                "strong", "beat", "exceed", "increase", "rise", "gain"
            ]
            
            bearish_keywords = [
                "sell", "bearish", "negative", "decline", "underperform", "downgrade",
                "weak", "miss", "fall", "decrease", "drop", "loss"
            ]
            
            content_lower = content.lower()
            
            bullish_signals = [word for word in bullish_keywords if word in content_lower]
            bearish_signals = [word for word in bearish_keywords if word in content_lower]
            
            # Calculate sentiment score
            sentiment_score = len(bullish_signals) - len(bearish_signals)
            
            if sentiment_score > 0:
                overall_sentiment = "Bullish"
            elif sentiment_score < 0:
                overall_sentiment = "Bearish"
            else:
                overall_sentiment = "Neutral"
            
            return {
                "success": True,
                "bullish_signals": bullish_signals,
                "bearish_signals": bearish_signals,
                "sentiment_score": sentiment_score,
                "overall_sentiment": overall_sentiment,
                "confidence": min(abs(sentiment_score) * 2, 10)  # Scale to 1-10
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}

    def analyze_and_store(self, document_id: str, symbol: Optional[str] = None) -> Dict:
        """Run AI analysis + signal extraction and persist results into Supabase for RAG."""
        try:
            if not self.supabase:
                return {"success": False, "error": "Supabase not configured"}

            # Retrieve content and ensure embedding present
            content = self.get_document_content(document_id) or ""
            embedding_vector = self._generate_embedding(content)

            # AI analysis and signals
            analysis = self.analyze_document_with_ai(document_id, symbol)
            signals = self.extract_trading_signals(document_id)

            # Prepare update payload with optional fields guarded
            update_payload = {
                "rag_embedding": embedding_vector if embedding_vector is not None else None,
                "last_analyzed_at": datetime.now().isoformat(),
            }

            if analysis.get("success"):
                update_payload.update({
                    "ai_analysis_text": analysis.get("analysis"),
                    # Simple extractions for recommendation/sentiment/confidence can be post-processed by ai_analysis module later
                })
            if signals.get("success"):
                update_payload.update({
                    "overall_sentiment": signals.get("overall_sentiment"),
                    "signal_confidence": signals.get("confidence"),
                    "bullish_signals": signals.get("bullish_signals"),
                    "bearish_signals": signals.get("bearish_signals"),
                })

            # Attempt update; tolerate missing columns
            try:
                self.supabase.table("research_documents").update(update_payload).eq("id", document_id).execute()
            except Exception as e:
                print(f"Supabase update warning (non-fatal): {e}")

            return {
                "success": True,
                "document_id": document_id,
                "analysis": analysis,
                "signals": signals,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def get_document_insights(self, symbol: str = None) -> List[Dict]:
        """Get AI insights for all documents or documents for a specific symbol"""
        try:
            documents = self.get_documents(symbol=symbol)
            insights = []
            
            for doc in documents:
                doc_id = doc.get("id")
                if doc_id:
                    # Get trading signals
                    signals = self.extract_trading_signals(doc_id)
                    
                    # Get AI analysis if available
                    analysis = self.analyze_document_with_ai(doc_id, symbol)
                    
                    insight = {
                        "document_id": doc_id,
                        "filename": doc.get("metadata", {}).get("file_name", "Unknown"),
                        "symbol": doc.get("metadata", {}).get("symbol", symbol),
                        "signals": signals,
                        "analysis": analysis,
                        "uploaded_at": doc.get("created_at", "Unknown")
                    }
                    insights.append(insight)
            
            return insights
            
        except Exception as e:
            print(f"Error getting document insights: {e}")
            return []
    
    def close(self):
        """Close database connection"""
        # No database connection to close for Supabase
        pass
